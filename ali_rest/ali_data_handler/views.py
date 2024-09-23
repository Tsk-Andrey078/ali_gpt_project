import os
from rest_framework.authentication import SessionAuthentication, BasicAuthentication
from rest_framework.permissions import AllowAny, IsAuthenticated
from rest_framework import status
from rest_framework.response import Response
from rest_framework.views import APIView
from rest_framework import viewsets
from django.contrib.auth import authenticate
from rest_framework.authtoken.models import Token
from django.contrib.auth.models import User
from openai import OpenAI
from docx import Document
from datetime import datetime, timedelta
from .models import CompanySell
from .serializer import CompanySellSerializer
from .clear_data import read_word_table

def find_vector_store_by_name(client, name):
    vector_stores = client.beta.vector_stores.list()
    for store in vector_stores:
        if store.name == name:
            return store
    return None

class UploadData(APIView):
    def post(self, request, *args, **kwargs):
        data = CompanySell.objects.all().values_list('data')
        doc = Document()
        doc.add_heading('Too_list', 0)

        for row in data:
            doc.add_paragraph(f"Описание компании: {row[0]}")
        
        # Сохранение документа в файл
        file_path = 'temp.docx'
        doc.save(file_path)

        # Готовим файл для загрузки в OpenAI
        file_stream = open(file_path, "rb")

        client = OpenAI(api_key="")

        vector_store = find_vector_store_by_name(client, "Too_List")
        if vector_store is None:
            vector_store = client.beta.vector_stores.create(name = "Too_List")
        
        file_batch = client.beta.vector_stores.file_batches.upload_and_poll(vector_store_id = vector_store.id, files = [file_stream])
        print(file_batch.status)
        print(file_batch.file_counts)
        print(file_batch.id)

        return Response(status = status.HTTP_200_OK)

class AddDataWord(APIView):
    def post(self, request, *args, **kwargs):
        uploadfile = request.FILES.get('company_list')
        print(uploadfile)
        if uploadfile:
            data = read_word_table(uploadfile)
       
            for i, row in data.iterrows():
                CompanySell.objects.create(data=row[0])
            
            return Response(status = status.HTTP_200_OK)
        return Response(status = status.HTTP_400_BAD_REQUEST)
    
class AddData(APIView):
    def post(self, request, *args, **kwargs):
        data = CompanySellSerializer(data = request.data)
        if data.is_valid():
            data.save()

            return Response(data.data, status=status.HTTP_201_CREATED)
        return Response(data.errors, status=status.HTTP_400_BAD_REQUEST)
    
